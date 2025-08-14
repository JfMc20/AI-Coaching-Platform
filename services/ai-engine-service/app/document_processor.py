"""
Document Processing Pipeline for AI Engine Service
Implements secure document processing with malware scanning, validation, and chunking
"""

import os
import logging
import asyncio
import hashlib
import magic
import tempfile
import shutil
from typing import List, Dict, Any, Optional, Tuple, BinaryIO
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass
from enum import Enum

# Document processing libraries with conditional imports
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None
    PYPDF2_AVAILABLE = False

try:
    import docx
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DocxDocument = None
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    markdown = None
    MARKDOWN_AVAILABLE = False

try:
    import magic
    # Test if libmagic is actually available
    magic.from_buffer(b"test", mime=True)
    MAGIC_AVAILABLE = True
except (ImportError, Exception):
    magic = None
    MAGIC_AVAILABLE = False

from shared.models.documents import DocumentChunk, ProcessingResult, ProcessingStatus
from shared.exceptions.documents import (
    DocumentProcessingError, UnsupportedFormatError, FileTooLargeError,
    MalwareDetectedError, TextExtractionError
)
from shared.ai.ollama_manager import get_ollama_manager, OllamaError
from shared.ai.chromadb_manager import get_chromadb_manager, ChromaDBError

logger = logging.getLogger(__name__)


class DocumentType(str, Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    MARKDOWN = "markdown"


@dataclass
class ProcessingConfig:
    """Configuration for document processing"""
    max_file_size_mb: int = 50
    chunk_size_tokens: int = 512
    chunk_overlap_tokens: int = 50
    max_chunks_per_document: int = 1000
    enable_malware_scan: bool = True
    enable_ocr: bool = False  # Future implementation
    quarantine_suspicious_files: bool = True


@dataclass
class SecurityScanResult:
    """Result of security scanning"""
    is_safe: bool
    threats_detected: List[str]
    scan_engine: str
    scan_time_ms: float


class SecurityScanner:
    """Handles security scanning of uploaded files"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.allowed_mime_types = {
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'text/plain',
            'text/markdown',
            'text/x-markdown'
        }
        self.dangerous_extensions = {
            '.exe', '.bat', '.cmd', '.com', '.pif', '.scr', '.vbs', '.js',
            '.jar', '.app', '.deb', '.pkg', '.dmg', '.iso', '.msi'
        }
    
    async def scan_file(
        self, 
        file_path: Path, 
        original_filename: str
    ) -> SecurityScanResult:
        """
        Perform comprehensive security scan on uploaded file
        
        Args:
            file_path: Path to temporary file
            original_filename: Original filename from upload
            
        Returns:
            Security scan result
            
        Raises:
            MalwareDetectedError: If malware is detected
            UnsupportedFormatError: If file type is not allowed
        """
        start_time = datetime.utcnow()
        threats = []
        
        try:
            # 1. Check file extension
            file_ext = Path(original_filename).suffix.lower()
            if file_ext in self.dangerous_extensions:
                threats.append(f"Dangerous file extension: {file_ext}")
            
            # 2. MIME type validation using python-magic
            try:
                if MAGIC_AVAILABLE:
                    detected_mime = magic.from_file(str(file_path), mime=True)
                    if detected_mime not in self.allowed_mime_types:
                        threats.append(f"Unsupported MIME type: {detected_mime}")
                else:
                    # Fallback MIME detection based on file extension
                    file_ext = Path(original_filename).suffix.lower()
                    mime_mapping = {
                        '.pdf': 'application/pdf',
                        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        '.txt': 'text/plain',
                        '.md': 'text/markdown'
                    }
                    detected_mime = mime_mapping.get(file_ext, 'application/octet-stream')
                    if detected_mime not in self.allowed_mime_types:
                        threats.append(f"Unsupported file type: {file_ext}")
            except Exception as e:
                logger.warning(f"MIME detection failed: {e}")
                threats.append("MIME type detection failed")
            
            # 3. File size check
            file_size = file_path.stat().st_size
            max_size_bytes = self.config.max_file_size_mb * 1024 * 1024
            if file_size > max_size_bytes:
                raise FileTooLargeError(
                    f"File size {file_size} bytes exceeds limit of {max_size_bytes} bytes"
                )
            
            # 4. Content-based malware patterns (basic)
            if await self._scan_content_patterns(file_path):
                threats.append("Suspicious content patterns detected")
            
            # 5. File structure validation
            if not await self._validate_file_structure(file_path, detected_mime):
                threats.append("Invalid file structure")
            
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            
            # Determine if file is safe
            is_safe = len(threats) == 0
            
            if not is_safe and self.config.quarantine_suspicious_files:
                await self._quarantine_file(file_path, original_filename, threats)
            
            if threats and any("malware" in threat.lower() for threat in threats):
                raise MalwareDetectedError(f"Malware detected: {', '.join(threats)}")
            
            if not is_safe:
                raise UnsupportedFormatError(f"File validation failed: {', '.join(threats)}")
            
            return SecurityScanResult(
                is_safe=is_safe,
                threats_detected=threats,
                scan_engine="internal",
                scan_time_ms=processing_time
            )
            
        except (MalwareDetectedError, UnsupportedFormatError, FileTooLargeError):
            raise
        except Exception as e:
            logger.error(f"Security scan failed: {e}")
            raise DocumentProcessingError(f"Security scan failed: {str(e)}") from e
    
    async def _scan_content_patterns(self, file_path: Path) -> bool:
        """
        Scan file content for suspicious patterns
        
        Args:
            file_path: Path to file
            
        Returns:
            True if suspicious patterns found
        """
        try:
            # Read first 1KB for pattern matching
            with open(file_path, 'rb') as f:
                content = f.read(1024)
            
            # Convert to string for pattern matching
            try:
                text_content = content.decode('utf-8', errors='ignore').lower()
            except Exception:
                return False
            
            # Suspicious patterns
            malicious_patterns = [
                b'<script',
                b'javascript:',
                b'vbscript:',
                b'eval(',
                b'exec(',
                b'system(',
                b'shell_exec(',
                b'<?php',
                b'<%',
                b'powershell',
                b'cmd.exe'
            ]
            
            for pattern in malicious_patterns:
                if pattern in content:
                    logger.warning(f"Suspicious pattern found: {pattern}")
                    return True
            
            return False
            
        except Exception as e:
            logger.warning(f"Content pattern scan failed: {e}")
            return False
    
    async def _validate_file_structure(self, file_path: Path, mime_type: str) -> bool:
        """
        Validate file structure matches expected format
        
        Args:
            file_path: Path to file
            mime_type: Detected MIME type
            
        Returns:
            True if structure is valid
        """
        try:
            if mime_type == 'application/pdf':
                # Basic PDF validation
                with open(file_path, 'rb') as f:
                    header = f.read(8)
                    return header.startswith(b'%PDF-')
            
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Basic DOCX validation (ZIP structure)
                try:
                    import zipfile
                    with zipfile.ZipFile(file_path, 'r') as zip_file:
                        # Check for required DOCX files
                        required_files = ['[Content_Types].xml', 'word/document.xml']
                        for required_file in required_files:
                            if required_file not in zip_file.namelist():
                                return False
                    return True
                except zipfile.BadZipFile:
                    return False
            
            elif mime_type in ['text/plain', 'text/markdown', 'text/x-markdown']:
                # Text files - check if readable
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    f.read(100)  # Try to read first 100 chars
                return True
            
            return True  # Default to valid for other types
            
        except Exception as e:
            logger.warning(f"File structure validation failed: {e}")
            return False
    
    async def _quarantine_file(
        self, 
        file_path: Path, 
        original_filename: str, 
        threats: List[str]
    ):
        """
        Quarantine suspicious file
        
        Args:
            file_path: Path to suspicious file
            original_filename: Original filename
            threats: List of detected threats
        """
        try:
            # Create quarantine directory
            quarantine_dir = Path(tempfile.gettempdir()) / "quarantine"
            quarantine_dir.mkdir(exist_ok=True)
            
            # Generate quarantine filename
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            quarantine_filename = f"{timestamp}_{original_filename}"
            quarantine_path = quarantine_dir / quarantine_filename
            
            # Move file to quarantine
            shutil.move(str(file_path), str(quarantine_path))
            
            # Log quarantine action
            logger.warning(
                f"File quarantined: {original_filename} -> {quarantine_path}, "
                f"Threats: {', '.join(threats)}"
            )
            
            # TODO: Send alert to administrators
            
        except Exception as e:
            logger.error(f"Failed to quarantine file: {e}")


class TextExtractor:
    """Extracts text from various document formats"""
    
    def __init__(self):
        self.extractors = {
            DocumentType.PDF: self._extract_pdf,
            DocumentType.DOCX: self._extract_docx,
            DocumentType.TXT: self._extract_txt,
            DocumentType.MD: self._extract_markdown,
            DocumentType.MARKDOWN: self._extract_markdown
        }
    
    async def extract_text(
        self, 
        file_path: Path, 
        document_type: DocumentType
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from document
        
        Args:
            file_path: Path to document file
            document_type: Type of document
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            TextExtractionError: If text extraction fails
        """
        try:
            if document_type not in self.extractors:
                raise UnsupportedFormatError(f"Unsupported document type: {document_type}")
            
            extractor = self.extractors[document_type]
            return await extractor(file_path)
            
        except Exception as e:
            logger.error(f"Text extraction failed for {document_type}: {e}")
            raise TextExtractionError(f"Failed to extract text: {str(e)}") from e
    
    async def _extract_pdf(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF file"""
        if not PYPDF2_AVAILABLE:
            raise TextExtractionError("PyPDF2 not available - PDF processing disabled in development mode")
        
        try:
            text_parts = []
            metadata = {"pages": 0, "title": "", "author": ""}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata["title"] = pdf_reader.metadata.get('/Title', '')
                    metadata["author"] = pdf_reader.metadata.get('/Author', '')
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
            
            full_text = "\n\n".join(text_parts)
            return full_text, metadata
            
        except Exception as e:
            raise TextExtractionError(f"PDF extraction failed: {str(e)}") from e
    
    async def _extract_docx(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise TextExtractionError("python-docx not available - DOCX processing disabled in development mode")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "paragraphs": len(doc.paragraphs)
            }
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            full_text = "\n\n".join(text_parts)
            return full_text, metadata
            
        except Exception as e:
            raise TextExtractionError(f"DOCX extraction failed: {str(e)}") from e
    
    async def _extract_txt(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            
            metadata = {
                "encoding": "utf-8",
                "lines": len(text.splitlines())
            }
            
            return text, metadata
            
        except Exception as e:
            raise TextExtractionError(f"TXT extraction failed: {str(e)}") from e
    
    async def _extract_markdown(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                markdown_text = file.read()
            
            if MARKDOWN_AVAILABLE:
                # Convert markdown to plain text (remove formatting)
                md = markdown.Markdown()
                html = md.convert(markdown_text)
                
                # Simple HTML tag removal
                import re
                text = re.sub(r'<[^>]+>', '', html)
                text = re.sub(r'\s+', ' ', text).strip()
            else:
                # Fallback: simple markdown processing
                import re
                text = markdown_text
                # Remove markdown formatting
                text = re.sub(r'#{1,6}\s+', '', text)  # Headers
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
                text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
                text = re.sub(r'`(.*?)`', r'\1', text)  # Code
                text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links
                text = re.sub(r'\s+', ' ', text).strip()
            
            metadata = {
                "format": "markdown",
                "lines": len(markdown_text.splitlines())
            }
            
            return text, metadata
            
        except Exception as e:
            raise TextExtractionError(f"Markdown extraction failed: {str(e)}") from e


class DocumentChunker:
    """Splits documents into chunks for embedding"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
    
    async def chunk_text(
        self, 
        text: str, 
        document_id: str,
        metadata: Dict[str, Any]
    ) -> List[DocumentChunk]:
        """
        Split text into chunks for embedding
        
        Args:
            text: Full document text
            document_id: Document identifier
            metadata: Document metadata
            
        Returns:
            List of document chunks
        """
        try:
            if not text.strip():
                return []
            
            # Simple sentence-based chunking with overlap
            sentences = self._split_into_sentences(text)
            chunks = []
            
            current_chunk = []
            current_tokens = 0
            chunk_index = 0
            
            for sentence in sentences:
                sentence_tokens = self._estimate_tokens(sentence)
                
                # If adding this sentence would exceed chunk size, create a new chunk
                if (current_tokens + sentence_tokens > self.config.chunk_size_tokens 
                    and current_chunk):
                    
                    # Create chunk
                    chunk_text = " ".join(current_chunk)
                    chunk = DocumentChunk(
                        id=f"{document_id}_chunk_{chunk_index}",
                        content=chunk_text,
                        metadata={
                            **metadata,
                            "document_id": document_id,
                            "chunk_index": chunk_index,
                            "source": "document_processor"
                        },
                        chunk_index=chunk_index,
                        token_count=current_tokens
                    )
                    chunks.append(chunk)
                    
                    # Start new chunk with overlap
                    overlap_sentences = self._get_overlap_sentences(
                        current_chunk, self.config.chunk_overlap_tokens
                    )
                    current_chunk = overlap_sentences + [sentence]
                    current_tokens = sum(self._estimate_tokens(s) for s in current_chunk)
                    chunk_index += 1
                else:
                    current_chunk.append(sentence)
                    current_tokens += sentence_tokens
                
                # Safety check for maximum chunks
                if len(chunks) >= self.config.max_chunks_per_document:
                    logger.warning(f"Reached maximum chunks limit for document {document_id}")
                    break
            
            # Add final chunk if there's remaining content
            if current_chunk:
                chunk_text = " ".join(current_chunk)
                chunk = DocumentChunk(
                    id=f"{document_id}_chunk_{chunk_index}",
                    content=chunk_text,
                    metadata={
                        **metadata,
                        "document_id": document_id,
                        "chunk_index": chunk_index,
                        "source": "document_processor"
                    },
                    chunk_index=chunk_index,
                    token_count=current_tokens
                )
                chunks.append(chunk)
            
            logger.info(f"Created {len(chunks)} chunks for document {document_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"Text chunking failed: {e}")
            raise DocumentProcessingError(f"Chunking failed: {str(e)}") from e
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        import re
        
        # Simple sentence splitting (can be improved with NLTK)
        sentences = re.split(r'[.!?]+', text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence) > 10:  # Minimum sentence length
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _estimate_tokens(self, text: str) -> int:
        """Estimate token count (rough approximation)"""
        # Simple approximation: 1 token ≈ 4 characters
        return len(text) // 4
    
    def _get_overlap_sentences(self, sentences: List[str], max_overlap_tokens: int) -> List[str]:
        """Get sentences for overlap between chunks"""
        overlap_sentences = []
        overlap_tokens = 0
        
        # Take sentences from the end for overlap
        for sentence in reversed(sentences):
            sentence_tokens = self._estimate_tokens(sentence)
            if overlap_tokens + sentence_tokens <= max_overlap_tokens:
                overlap_sentences.insert(0, sentence)
                overlap_tokens += sentence_tokens
            else:
                break
        
        return overlap_sentences


class DocumentProcessor:
    """
    Main document processor with security scanning, text extraction, and chunking
    """
    
    def __init__(self, config: ProcessingConfig = None):
        self.config = config or ProcessingConfig()
        self.security_scanner = SecurityScanner(self.config)
        self.text_extractor = TextExtractor()
        self.chunker = DocumentChunker(self.config)
        self.ollama_manager = get_ollama_manager()
        self.chromadb_manager = get_chromadb_manager()
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str,
        creator_id: str,
        document_id: Optional[str] = None
    ) -> ProcessingResult:
        """
        Process uploaded document through complete pipeline
        
        Args:
            file_content: Raw file content
            filename: Original filename
            creator_id: Creator identifier for tenant isolation
            document_id: Optional document ID (generated if not provided)
            
        Returns:
            Processing result with chunks and metadata
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        start_time = datetime.utcnow()
        document_id = document_id or self._generate_document_id(filename, creator_id)
        
        # Create temporary file
        temp_file = None
        try:
            # 1. Save to temporary file
            temp_file = await self._save_temp_file(file_content, filename)
            
            # 2. Security scanning
            logger.info(f"Starting security scan for document {document_id}")
            scan_result = await self.security_scanner.scan_file(temp_file, filename)
            
            # 3. Determine document type
            document_type = self._detect_document_type(filename)
            
            # 4. Extract text
            logger.info(f"Extracting text from document {document_id}")
            text, metadata = await self.text_extractor.extract_text(temp_file, document_type)
            
            if not text.strip():
                raise TextExtractionError("No text content extracted from document")
            
            # 5. Chunk text
            logger.info(f"Chunking text for document {document_id}")
            chunks = await self.chunker.chunk_text(text, document_id, metadata)
            
            if not chunks:
                raise DocumentProcessingError("No chunks created from document")
            
            # 6. Generate embeddings
            logger.info(f"Generating embeddings for {len(chunks)} chunks")
            await self._generate_embeddings(chunks, creator_id)
            
            # 7. Store in ChromaDB
            logger.info(f"Storing chunks in ChromaDB for creator {creator_id}")
            await self._store_chunks(chunks, creator_id)
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # 8. Create processing result
            result = ProcessingResult(
                id=f"result_{document_id}",
                creator_id=creator_id,
                document_id=document_id,
                status=ProcessingStatus.COMPLETED,
                chunks=chunks,
                total_chunks=len(chunks),
                processing_time_seconds=processing_time,
                metadata={
                    "filename": filename,
                    "document_type": document_type.value,
                    "security_scan": {
                        "is_safe": scan_result.is_safe,
                        "scan_engine": scan_result.scan_engine,
                        "scan_time_ms": scan_result.scan_time_ms
                    },
                    "text_extraction": metadata,
                    "chunking": {
                        "chunk_size_tokens": self.config.chunk_size_tokens,
                        "chunk_overlap_tokens": self.config.chunk_overlap_tokens
                    }
                }
            )
            
            logger.info(
                f"Document processing completed for {document_id}: "
                f"{len(chunks)} chunks in {processing_time:.2f}s"
            )
            
            return result
            
        except Exception as e:
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # Create failed result
            error_result = ProcessingResult(
                id=f"result_{document_id}",
                creator_id=creator_id,
                document_id=document_id,
                status=ProcessingStatus.FAILED,
                chunks=[],
                total_chunks=0,
                processing_time_seconds=processing_time,
                error_message=str(e),
                metadata={
                    "filename": filename,
                    "error_type": type(e).__name__
                }
            )
            
            logger.error(f"Document processing failed for {document_id}: {e}")
            return error_result
            
        finally:
            # Cleanup temporary file
            if temp_file and temp_file.exists():
                try:
                    temp_file.unlink()
                except Exception as e:
                    logger.warning(f"Failed to cleanup temp file: {e}")
    
    async def _save_temp_file(self, file_content: bytes, filename: str) -> Path:
        """Save uploaded content to temporary file"""
        try:
            # Create temporary file with original extension
            suffix = Path(filename).suffix
            temp_file = Path(tempfile.mktemp(suffix=suffix))
            
            with open(temp_file, 'wb') as f:
                f.write(file_content)
            
            return temp_file
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to save temporary file: {str(e)}") from e
    
    def _detect_document_type(self, filename: str) -> DocumentType:
        """Detect document type from filename"""
        extension = Path(filename).suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MD,
            '.markdown': DocumentType.MARKDOWN
        }
        
        if extension not in type_mapping:
            raise UnsupportedFormatError(f"Unsupported file extension: {extension}")
        
        return type_mapping[extension]
    
    def _generate_document_id(self, filename: str, creator_id: str) -> str:
        """Generate unique document ID"""
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename_hash = hashlib.md5(filename.encode()).hexdigest()[:8]
        return f"doc_{creator_id}_{timestamp}_{filename_hash}"
    
    async def _generate_embeddings(self, chunks: List[DocumentChunk], creator_id: str):
        """Generate embeddings for document chunks"""
        try:
            # Extract text content from chunks
            texts = [chunk.content for chunk in chunks]
            
            # Generate embeddings in batches
            batch_size = 10
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                batch_chunks = chunks[i:i + batch_size]
                
                # Generate embeddings for batch
                embedding_response = await self.ollama_manager.generate_embeddings(batch_texts)
                
                # Assign embeddings to chunks
                for j, embedding in enumerate(embedding_response.embeddings):
                    if i + j < len(chunks):
                        batch_chunks[j].embedding_vector = embedding
            
            logger.info(f"Generated embeddings for {len(chunks)} chunks")
            
        except OllamaError as e:
            raise DocumentProcessingError(f"Embedding generation failed: {str(e)}") from e
    
    async def _store_chunks(self, chunks: List[DocumentChunk], creator_id: str):
        """Store document chunks in ChromaDB"""
        try:
            # Prepare data for ChromaDB
            embeddings = [chunk.embedding_vector for chunk in chunks]
            documents = [chunk.content for chunk in chunks]
            metadatas = [chunk.metadata for chunk in chunks]
            ids = [chunk.id for chunk in chunks]
            
            # Store in ChromaDB
            await self.chromadb_manager.add_embeddings(
                creator_id=creator_id,
                document_id=chunks[0].metadata["document_id"],
                embeddings=embeddings,
                documents=documents,
                metadatas=metadatas,
                ids=ids
            )
            
            logger.info(f"Stored {len(chunks)} chunks in ChromaDB for creator {creator_id}")
            
        except ChromaDBError as e:
            raise DocumentProcessingError(f"ChromaDB storage failed: {str(e)}") from e


# Global document processor instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get global document processor instance (lazy initialization)"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor